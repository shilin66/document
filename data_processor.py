#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
核心数据处理器模块
专注于生成模板变量，保持源文件格式
"""
import copy
import io
import os
from datetime import datetime, timedelta
from typing import Dict, List, Any
from docx import Document

try:
    from excel2img import process_file
except ImportError:
    print("警告：excel2img模块导入失败，Excel转图片功能将不可用")
    def process_file(*args, **kwargs):
        print("Excel转图片功能不可用")
        return []

try:
    from pdf_parser import PDFParser
except ImportError:
    print("警告：pdf_parser模块导入失败，PDF解析功能将不可用")
    PDFParser = None
from file_scanner import FileScanner
from minio_file_scanner import MinioFileScanner


class CoreDataProcessor:
    """核心数据处理器类"""

    def __init__(self, use_minio=False, minio_config=None, config=None):
        """初始化数据处理器
        
        Args:
            use_minio: 是否使用Minio作为文件源
            minio_config: Minio配置字典
            config: 完整的配置对象
        """
        self.use_minio = use_minio
        self.config = config
        
        # 初始化PDF解析器
        self.pdf_parser = None
        if PDFParser:
            pdf_api_url = config.get_pdf_api_url() if config else 'http://187.9.9.8:7434/v2/parse/file'
            pdf_parse_mode = config.get_pdf_parse_mode() if config else 'api'
            pdf_image_dpi = config.get_pdf_to_image_dpi() if config else 200
            poppler_path = config.get_poppler_path() if config else None
            
            self.pdf_parser = PDFParser(
                api_url=pdf_api_url,
                parse_mode=pdf_parse_mode,
                image_dpi=pdf_image_dpi,
                poppler_path=poppler_path
            )
        
        if use_minio and minio_config:
            # 使用配置中的base_prefix
            base_prefix = minio_config.get('base_prefix', '核心网络部运维报告')
            self.scanner = MinioFileScanner(minio_config, base_prefix=base_prefix)
        else:
            # 使用本地配置
            local_config = config.get_local_config() if config else {}
            base_path = local_config.get('base_path', '核心网络部运维报告')
            self.scanner = FileScanner(base_path)

        # 定义变量映射关系
        self.variable_mapping = {
            'report_date': '报告生成时间',
            'year': '年份',
            'month': '月份',
            'network_failure_stats': '网络故障统计',
            'blackhole_ip_stats': '黑洞路由器IP统计',
            'ip_statistics': 'IP统计信息',
            'fault_drill_plan_unicom_night': '中国联通出口故障演习方案（夜班及节假日）',
            'fault_drill_plan_unicom_day': '中国联通出口故障演习方案（白班）',
            'fault_drill_plan_unicom_international_night': '联通国际出口故障演习方案（夜班及节假日）',
            'fault_drill_plan_unicom_international_day': '联通国际出口故障演习方案（白班）',
            'internet_traffic_stats': '互联网出口业务流量统计',
            'isp_bandwidth': 'ISP业务宽带复用比',
            'idc_traffic_stats': 'IDC业务出口流量统计',
            'isp_traffic_stats': 'ISP业务出口流量统计',
            'dedicated_traffic_peak': '专线托管峰值流量',
            'bandwidth_usage': '带宽使用统计',
            'core_device_hardware': '核心设备硬件指标',
            'monthly_work_plan': '计划性项目',
            'sdwan_project': 'SDWAN项目',
            'device_running_report_and_suggestion':'网络设备运行报告和建议',
        }

    def process_all_files(self, template_variables: set = None, target_date: str = None) -> Dict[str, str]:
        """
        处理所有文件，生成模板变量字典

        Args:
            template_variables: 模板中存在的变量集合，如果为None则处理所有变量
            target_date: 目标日期，格式为YYYYMM。如果为None，则使用上个月

        Returns:
            Dict[str, str]: 模板变量字典
        """
        # 初始化变量字典
        variables = {}

        # 如果指定了模板变量，只初始化这些变量
        if template_variables:
            for var_name in template_variables:
                if var_name in self.variable_mapping:
                    variables[var_name] = f"[{self.variable_mapping[var_name]}：暂无数据]"
                else:
                    variables[var_name] = f"[未知变量: {var_name}]"
        else:
            # 处理所有变量
            for var_name in self.variable_mapping.keys():
                variables[var_name] = f"[{self.variable_mapping[var_name]}：暂无数据]"

        # 设置报告生成时间
        if 'report_date' in variables:
            variables['report_date'] = datetime.now().strftime('%Y年%m月%d日')

        if target_date is None or (isinstance(target_date, str) and target_date.strip() == ""):
            # 获取上个月的年月
            today = datetime.today()
            first_day_current_month = today.replace(day=1)
            last_month = first_day_current_month - timedelta(days=1)
            target_date = last_month.strftime("%Y%m")
            variables['year'] = last_month.strftime("%Y")
            variables['month'] = last_month.strftime("%m")
        else:
            target_date = target_date.strip()
            variables['year'] = target_date[:4]
            variables['month'] = target_date[4:]
        # 扫描所有文件
        if self.use_minio:
            # 如果使用Minio，则传递target_date参数
            file_dict = self.scanner.scan_files(target_date)
        else:
            # 如果使用本地文件，则忽略target_date参数
            file_dict = self.scanner.scan_files()

        # 处理各个目录的文件
        for directory, files in file_dict.items():
            if "监控中心运维报告" in directory:
                self._process_monitor_center_operation_report_files(files, variables)
            elif "核心网络流量统计报告" in directory:
                self._process_core_network_traffic_statistics_report_files(files, variables)
            elif "核心设备运行报告" in directory:
                self._process_core_equipment_operation_report_files(files, variables)
            elif "月度计划及总结" in directory:
                self._process_monthly_plan_and_summary_files(files, variables)

        return variables

    def _process_monitor_center_operation_report_files(self, files: List[str], variables: Dict[str, Any]):
        """处理张嵩的运维报告文件，保持完整格式"""
        summary_parts = []

        for file_path in files:
            filename = os.path.basename(file_path)

            if file_path.endswith('.pdf'):
                # 处理PDF文件
                self._process_pdf_file(file_path, variables)
            elif file_path.endswith('.docx'):
                # content = self.docx_reader.read_docx_content(file_path)

                if '网络故障统计' in filename:
                    if 'network_failure_stats' in variables:
                        # 使用完整格式信息提取第一个表格
                        variables['network_failure_stats'] = {
                            'type': 'sub_doc',
                            'value': self._extract_first_table_from_docx(file_path)
                        }
                        summary_parts.append("网络故障统计报告")

                elif '黑洞路由器' in filename:
                    if 'blackhole_ip_stats' in variables:
                        variables['blackhole_ip_stats'] = {
                            'type': 'sub_doc',
                            'value': file_path
                        }
                        summary_parts.append("黑洞路由器IP统计")

                elif 'IP' in filename:
                    if 'ip_statistics' in variables:
                        variables['ip_statistics'] = {
                            'type': 'sub_doc',
                            'value': file_path
                        }
                        summary_parts.append("IP统计信息")

            elif file_path.endswith('.xlsx'):
                # 验证文件是否存在
                if not os.path.exists(file_path):
                    print(f"警告：文件不存在: {file_path}")
                    continue
                    
                if '中国联通1出口故障演习' in filename:
                    poppler_path = self.config.get_poppler_path() if self.config else None
                    if '白班' in filename:
                        imgs = process_file(file_path, 200, None, poppler_path, False, '演练记录')
                        if len(imgs) > 0:
                            if  len(imgs) == 1:
                                variables['fault_drill_plan_unicom_day'] = {
                                    'type': 'image',
                                    'value': imgs[0]
                                }
                            else:
                                variables['fault_drill_plan_unicom_day'] = {
                                    'type': 'image',
                                    'value': imgs[1]
                                }
                    if '夜班' in filename:
                        imgs = process_file(file_path, 200, None, poppler_path, False, '演练记录')
                        if len(imgs) > 0:
                            if len(imgs) == 1:
                                variables['fault_drill_plan_unicom_night'] = {
                                    'type': 'image',
                                    'value': imgs[0]
                                }
                            else:
                                variables['fault_drill_plan_unicom_night'] = {
                                    'type': 'image',
                                    'value': imgs[1]
                                }
                if '联通国际出口' in filename:
                    poppler_path = self.config.get_poppler_path() if self.config else None
                    if '白班' in filename:
                        imgs = process_file(file_path, 200, None, poppler_path, False, '演练记录')
                        if len(imgs) > 0:
                            if len(imgs) == 1:
                                variables['fault_drill_plan_unicom_international_day'] = {
                                    'type': 'image',
                                    'value': imgs[0]
                                }
                            else:
                                variables['fault_drill_plan_unicom_international_day'] = {
                                    'type': 'image',
                                    'value': imgs[1]
                                }
                    if '夜班' in filename:
                        imgs = process_file(file_path, 200, None, poppler_path, False, '演练记录')
                        if len(imgs) > 0:
                            if len(imgs) == 1:
                                variables['fault_drill_plan_unicom_international_night'] = {
                                    'type': 'image',
                                    'value': imgs[0]
                                }
                            else:
                                variables['fault_drill_plan_unicom_international_night'] = {
                                    'type': 'image',
                                    'value': imgs[1]
                                }
    def _process_core_network_traffic_statistics_report_files(self, files: List[str], variables: Dict[str, Any]):
        """处理李曦炎的运维报告文件，保持完整格式"""
        for file_path in files:
            filename = os.path.basename(file_path)

            if file_path.endswith('.pdf'):
                # 处理PDF文件
                self._process_pdf_file(file_path, variables)
            elif file_path.endswith('.docx'):
                if '专线托管' in filename:
                    if 'dedicated_traffic_peak' in variables:
                        variables['dedicated_traffic_peak'] = {
                            'type': 'sub_doc',
                            'value': file_path
                        }

                elif '带宽使用' in filename:
                    if 'bandwidth_usage' in variables:
                        variables['bandwidth_usage'] = {
                            'type': 'sub_doc',
                            'value': file_path
                        }

            elif file_path.endswith('.xlsx'):
                # 验证文件是否存在
                if not os.path.exists(file_path):
                    print(f"警告：文件不存在: {file_path}")
                    continue
                    
                if '互联网出口业务流量' in filename:
                    poppler_path = self.config.get_poppler_path() if self.config else None
                    imgs = process_file(file_path, 200, None, poppler_path, False, None )
                    if len(imgs) > 0:
                        variables['internet_traffic_stats'] = {
                            'type': 'image',
                            'value': imgs[0]
                        }
                        if imgs[1]:
                            variables['isp_bandwidth'] = {
                                'type': 'image',
                                'value': imgs[1]
                            }

                elif 'IDC业务出口流量' in filename:
                    poppler_path = self.config.get_poppler_path() if self.config else None
                    imgs = process_file(file_path, 200, None, poppler_path, False, None)
                    if len(imgs) > 0:
                        variables['idc_traffic_stats'] = {
                            'type': 'image',
                            'value': imgs[0]
                        }
                elif 'ISP业务出口流量' in filename:
                    poppler_path = self.config.get_poppler_path() if self.config else None
                    imgs = process_file(file_path, 200, None, poppler_path, False, None)
                    if len(imgs) > 0:
                        variables['isp_traffic_stats'] = {
                            'type': 'image',
                            'value': imgs[0]
                        }

    def _process_core_equipment_operation_report_files(self, files: List[str], variables: Dict[str, Any]):
        """处理王子徽的监控报告文件"""
        summary_parts = []

        for file_path in files:
            filename = os.path.basename(file_path)

            if file_path.endswith('.pdf'):
                # 处理PDF文件
                self._process_pdf_file(file_path, variables)
            elif '核心设备硬件指标' in filename:
                if 'core_device_hardware' in variables:
                    variables['core_device_hardware'] = {
                        'type': 'image',
                        'value': f"核心设备硬件指标报告：{filename}\n文件路径：{file_path}"
                    }
                    summary_parts.append("核心设备硬件指标")

    def _process_monthly_plan_and_summary_files(self, files: List[str], variables: Dict[str, Any]):
        """处理陈斌的月度计划总结文件"""

        for file_path in files:
            filename = os.path.basename(file_path)

            if file_path.endswith('.pdf'):
                # 处理PDF文件
                self._process_pdf_file(file_path, variables)
            elif file_path.endswith('.docx'):
                if '报告及建议' in filename:
                    if 'device_running_report_and_suggestion' in variables:
                        variables['device_running_report_and_suggestion'] = {
                            'type': 'sub_doc',
                            'value': file_path
                        }
            elif file_path.endswith('.xlsx'):
                # 验证文件是否存在
                if not os.path.exists(file_path):
                    print(f"警告：文件不存在: {file_path}")
                    continue
                    
                poppler_path = self.config.get_poppler_path() if self.config else None
                imgs = process_file(file_path, 200, None, poppler_path, False, None)
                if len(imgs) > 0:
                    variables['monthly_work_plan'] = {
                        'type': 'image',
                        'value': imgs[0]
                    }
                    if imgs[1]:
                        variables['sdwan_project'] = {
                            'type': 'image',
                            'value': imgs[1]
                        }

    def generate_appendix(self, file_dict: Dict[str, List[str]]) -> str:
        """生成附录内容"""
        appendix_parts = []

        appendix_parts.append("附录：源文件清单")
        appendix_parts.append("=" * 30)

        for directory, files in file_dict.items():
            appendix_parts.append(f"\n{directory}:")
            for file_path in files:
                filename = os.path.basename(file_path)
                appendix_parts.append(f"  - {filename}")

        appendix_parts.append(f"\n总计文件数量：{sum(len(files) for files in file_dict.values())}")
        appendix_parts.append(f"处理时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        return '\n'.join(appendix_parts)

    def _extract_first_table_from_docx(self, file_path: str):
        """
        从 DOCX 内容中提取第一个表格的数据，保持完整格式

        Args:
            content: 由 CoreDocxReader.read_docx_content 返回的内容字典

        Returns:
            str: 第一个表格的数据，带有完整格式信息或管道分隔符格式
        """
        doc = Document(file_path)
        first_table = copy.deepcopy(doc.tables[0]._element)
        new_doc = Document()
        new_doc.element.body.append(first_table)
        bio = io.BytesIO()
        new_doc.save(bio)
        bio.seek(0)
        return bio

    def _extract_formatted_content(self, file_path: str) -> str:
        """
        提取DOCX内容的完整格式信息，保持所有表格和段落的格式

        Args:
            content: 由 CoreDocxReader.read_docx_content 返回的内容字典

        Returns:
            str: 带有完整格式信息的内容
        """
        doc = Document(file_path)
        result_parts = doc._element.body

        return result_parts

    def _process_pdf_file(self, file_path: str, variables: Dict[str, Any]):
        """
        处理PDF文件
        
        Args:
            file_path: PDF文件路径
            variables: 变量字典
        """
        if not self.pdf_parser:
            print(f"警告：PDF解析器不可用，跳过文件: {file_path}")
            return
        
        try:
            print(f"正在处理PDF文件: {os.path.basename(file_path)}")
            
            # 创建图片输出目录
            from datetime import datetime
            from temp_utils import create_pdf_images_temp_dir
            date_str = datetime.now().strftime('%Y%m%d')
            image_dir = create_pdf_images_temp_dir()
            
            # 解析PDF
            result = self.pdf_parser.parse_pdf(
                file_path, 
                extract_images=True,
                image_output_dir=image_dir
            )
            
            if result['success']:
                # 检查解析模式
                if result.get('parse_mode') == 'image':
                    # 图片模式：直接使用图片路径数组
                    image_paths = result.get('image_paths', [])
                    
                    # 如果有老的变量名，也处理一下
                    if 'core_device_hardware' in variables:
                        variables['core_device_hardware'] = {
                            'type': 'image_array',
                            'value': image_paths
                        }
                else:
                    # API模式：使用文本内容
                    markdown_content = result['markdown']

                    # 如果有老的变量名，也处理一下
                    if 'core_device_hardware' in variables:
                        variables['core_device_hardware'] = {
                            'type': 'text',
                            'value': markdown_content
                        }
                
                print(f"✓ PDF处理成功: {os.path.basename(file_path)}")
                if result.get('parse_mode') == 'image':
                    print(f"  生成图片: {len(result.get('image_paths', []))} 张")
                else:
                    print(f"  内容长度: {len(result.get('markdown', ''))} 字符")
                    print(f"  图片数量: {len(result.get('images', []))}")
                
            else:
                error_msg = result.get('error', '未知错误')
                print(f"✗ PDF处理失败: {error_msg}")
                
                # 设置错误信息
                if 'core_device_hardware' in variables:
                    variables['core_device_hardware'] = f"[PDF解析失败: {error_msg}]"
        
        except Exception as e:
            print(f"✗ PDF处理异常: {e}")
            
            # 设置异常信息
            if 'core_device_hardware' in variables:
                variables['core_device_hardware'] = f"[PDF处理异常: {str(e)}]"